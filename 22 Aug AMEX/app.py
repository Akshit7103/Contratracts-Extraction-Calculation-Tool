import os
import pandas as pd
from flask import Flask, render_template, request, redirect, url_for, flash
from docx import Document
import re

def extract_chd_rules_from_docx(docx_path: str):
    """Extract CHD (Client Held Day) rules from Word contract"""
    doc = Document(docx_path)
    chd_rules = {
        "threshold_days": None,
        "adjustment_bps_per_day": None,
        "description": "",
        "adjustment_type": None  # "deduct" or "add"
    }
    
    # Search for CHD rules in paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        text_lower = text.lower()
        
        # Look for CHD-related content
        if 'chd' in text_lower and any(keyword in text_lower for keyword in ['days', 'bps', 'adjustment', 'performance']):
            
            # Look for threshold patterns (e.g., "14 days or more", "is 14 days or more")
            threshold_patterns = [
                r'(\d+(?:\.\d+)?)\s*days?\s+or\s+more',
                r'is\s+(\d+(?:\.\d+)?)\s*days?\s+or\s+more',
                r'(\d+(?:\.\d+)?)\s*days?\s+or\s+greater',
                r'>=?\s*(\d+(?:\.\d+)?)\s*days?',
                r'(\d+(?:\.\d+)?)\s*days?\s+or\s+above'
            ]
            
            for pattern in threshold_patterns:
                match = re.search(pattern, text_lower)
                if match:
                    chd_rules["threshold_days"] = float(match.group(1))
                    break
            
            # Look for bps adjustment patterns (e.g., "2 bps per day", "2.5 bps per day")
            bps_patterns = [
                r'(\d+(?:\.\d+)?)\s*bps?\s+per\s+day',
                r'(\d+(?:\.\d+)?)\s*basis\s+points?\s+per\s+day',
                r'(\d+(?:\.\d+)?)\s*bp\s+per\s+day'
            ]
            
            for pattern in bps_patterns:
                match = re.search(pattern, text_lower)
                if match:
                    chd_rules["adjustment_bps_per_day"] = float(match.group(1))
                    break
            
            # Determine if it's a deduction or addition
            if any(word in text_lower for word in ['deduct', 'subtract', 'reduce', 'penalty']):
                chd_rules["adjustment_type"] = "deduct"
            elif any(word in text_lower for word in ['add', 'bonus', 'increase']):
                chd_rules["adjustment_type"] = "add"
            elif 'will be' in text_lower and any(word in text_lower for word in ['deduct', 'subtract']):
                chd_rules["adjustment_type"] = "deduct"
                
            # Store the description (keep the most detailed one)
            if len(text) > len(chd_rules["description"]):
                chd_rules["description"] = text
    
    # If we found both threshold and bps but no adjustment type, assume deduction (most common)
    if chd_rules["threshold_days"] is not None and chd_rules["adjustment_bps_per_day"] is not None and chd_rules["adjustment_type"] is None:
        chd_rules["adjustment_type"] = "deduct"
    
    return chd_rules

def extract_tiers_from_docx(docx_path: str):
    """Extract tier rules from Word contract"""
    doc = Document(docx_path)
    tiers = []
    
    def parse_range_cell(cell_text):
        t = " ".join(cell_text.split())
        t = t.replace("–", "-").replace("—", "-").replace("−", "-")
        if "+" in t:
            lower = re.findall(r"[-+]?\d*\.?\d+", t)
            if lower:
                return float(lower[0]), None
        parts = re.findall(r"[-+]?\d*\.?\d+", t)
        if len(parts) >= 2:
            return float(parts[0]), float(parts[1])
        return None, None

    def parse_bps_cell(cell_text):
        nums = re.findall(r"[-+]?\d*\.?\d+", cell_text)
        if nums:
            return float(nums[0])
        return None

    for tbl in doc.tables:
        header_text = " ".join([cell.text for cell in tbl.rows[0].cells]).lower()
        if ("nacv" in header_text and "bp" in header_text) or ("bp multiplier" in header_text):
            for i, row in enumerate(tbl.rows[1:], start=1):
                cells = [c.text.strip() for c in row.cells]
                if len(cells) < 2:
                    continue
                rng_text = cells[0]
                bps_text = cells[-1]
                lo, hi = parse_range_cell(rng_text)
                bps = parse_bps_cell(bps_text)
                if lo is not None and bps is not None:
                    tiers.append({
                        "lower_m": lo,
                        "upper_m": hi,
                        "bps": bps
                    })
    
    # Remove duplicates and sort
    dedup = []
    seen = set()
    for t in tiers:
        key = (round(t["lower_m"], 2), None if t["upper_m"] is None else round(t["upper_m"], 2), round(t["bps"], 2))
        if key not in seen:
            seen.add(key)
            dedup.append(t)
    
    dedup.sort(key=lambda x: (x["lower_m"], x["upper_m"] if x["upper_m"] is not None else 10**9))
    return dedup

def compute_chd_adjustment(gross_incentive_gbp: float, chd_performance_days: float, chd_rules):
    """Compute CHD adjustment to gross incentive"""
    if not chd_rules or chd_rules["threshold_days"] is None or chd_rules["adjustment_bps_per_day"] is None:
        return 0.0, {"applied": False, "reason": "No CHD rules found"}
    
    if chd_performance_days < chd_rules["threshold_days"]:
        return 0.0, {
            "applied": False, 
            "reason": f"CHD performance ({chd_performance_days} days) below threshold ({chd_rules['threshold_days']} days)"
        }
    
    # Calculate adjustment: bps per day based on contract
    adjustment_bps = chd_performance_days * chd_rules["adjustment_bps_per_day"]
    adjustment_rate = adjustment_bps / 10000.0
    adjustment_amount = gross_incentive_gbp * adjustment_rate
    
    # Determine the adjustment type
    adjustment_type = chd_rules.get("adjustment_type", "deduct")
    
    return adjustment_amount, {
        "applied": True,
        "chd_performance_days": chd_performance_days,
        "threshold_days": chd_rules["threshold_days"],
        "adjustment_bps_per_day": chd_rules["adjustment_bps_per_day"],
        "adjustment_type": adjustment_type,
        "total_adjustment_bps": adjustment_bps,
        "adjustment_amount_gbp": adjustment_amount
    }

def compute_gross_incentive(nacv_value_gbp: float, tiers):
    """Compute gross incentive using tier rules"""
    if not tiers:
        raise ValueError("No tiers found in the contract.")

    # Find top tier (with upper_m = None)
    top_idx = None
    for i, t in enumerate(tiers):
        if t["upper_m"] is None:
            top_idx = i
            break
    if top_idx is None:
        top_idx = len(tiers) - 1

    top = tiers[top_idx]
    if top_idx == 0:
        raise ValueError("Tier list malformed (no lower tiers before top).")

    second = tiers[top_idx - 1]
    second_upper_m = second["upper_m"] if second["upper_m"] is not None else second["lower_m"]

    nacv_m = nacv_value_gbp / 1_000_000.0

    if nacv_m < top["lower_m"]:
        # Single bracket calculation
        chosen = None
        for t in tiers:
            lo = t["lower_m"]
            hi = t["upper_m"]
            if hi is None:
                continue
            if nacv_m >= lo and nacv_m <= hi:
                chosen = t
                break
        if chosen is None:
            chosen = tiers[0]
        
        bps = chosen["bps"]
        rate = bps / 10000.0
        gross = nacv_value_gbp * rate
        breakdown = {
            "mode": "single-bracket",
            "bracket_bps": bps,
            "applied_amount_gbp": nacv_value_gbp,
            "calculation_steps": [
                f"NACV Value: {format_currency(nacv_value_gbp)}",
                f"Applied Rate: {bps:.1f} bps ({rate:.4%})",
                f"Calculation: {format_currency(nacv_value_gbp)} × {rate:.4%} = {format_currency(gross)}"
            ],
            "explanation": f"Since your NACV of {format_currency(nacv_value_gbp)} falls within the {chosen['lower_m']:.2f}M - {chosen['upper_m']:.2f}M tier, we apply a single rate of {bps:.1f} bps to the entire amount."
        }
        return gross, breakdown
    else:
        # Two-tier calculation
        second_upper_gbp = second_upper_m * 1_000_000.0
        top_bps = top["bps"]
        second_bps = second["bps"]
        top_rate = top_bps / 10000.0
        second_rate = second_bps / 10000.0

        base_amount = min(nacv_value_gbp, second_upper_gbp)
        remainder = max(0.0, nacv_value_gbp - second_upper_gbp)

        base_incentive = base_amount * second_rate
        remainder_incentive = remainder * top_rate
        gross = base_incentive + remainder_incentive
        
        breakdown = {
            "mode": "two-tier",
            "second_upper_m": second_upper_m,
            "second_bps": second_bps,
            "top_bps": top_bps,
            "base_amount_gbp": base_amount,
            "remainder_amount_gbp": remainder,
            "base_incentive": base_incentive,
            "remainder_incentive": remainder_incentive,
            "calculation_steps": [
                f"NACV Value: {format_currency(nacv_value_gbp)}",
                f"Tier 1 (up to {second_upper_m:.2f}M): {format_currency(base_amount)} × {second_bps:.1f} bps ({second_rate:.4%}) = {format_currency(base_incentive)}",
                f"Tier 2 (above {second_upper_m:.2f}M): {format_currency(remainder)} × {top_bps:.1f} bps ({top_rate:.4%}) = {format_currency(remainder_incentive)}",
                f"Total Gross Incentive: {format_currency(base_incentive)} + {format_currency(remainder_incentive)} = {format_currency(gross)}"
            ],
            "explanation": f"Your NACV of {format_currency(nacv_value_gbp)} exceeds {second_upper_m:.2f}M, so we use a two-tier calculation: the first {format_currency(second_upper_gbp)} at {second_bps:.1f} bps, and the remaining {format_currency(remainder)} at {top_bps:.1f} bps."
        }
        return gross, breakdown

def calculate_monthly_totals(excel_path: str):
    """
    Simple function to calculate totals from monthly data.
    Expects horizontal format: rows = categories, columns = months
    """
    df = pd.read_excel(excel_path, engine="openpyxl")
    
    # Find the rows for each category
    bus_row = None
    ar_row = None
    writeoff_row = None
    
    for i, row_label in enumerate(df.iloc[:, 0]):
        if pd.isna(row_label):
            continue
        label_lower = str(row_label).lower()
        if "bus unadjusted" in label_lower:
            bus_row = i
        elif "ar 180" in label_lower:
            ar_row = i
        elif "writeoff" in label_lower and "reserve" in label_lower:
            writeoff_row = i
    
    if bus_row is None or ar_row is None or writeoff_row is None:
        raise ValueError("Could not find required data rows in Excel file")
    
    # First, calculate totals from individual month columns
    numeric_cols = []
    for col_idx in range(1, len(df.columns)):  # Skip first column (labels)
        if pd.api.types.is_numeric_dtype(df.iloc[:, col_idx]):
            numeric_cols.append(col_idx)
    
    bus_total_calculated = 0.0
    ar_total_calculated = 0.0
    writeoff_total_calculated = 0.0
    
    for col_idx in numeric_cols:
        bus_val = df.iloc[bus_row, col_idx]
        ar_val = df.iloc[ar_row, col_idx]
        writeoff_val = df.iloc[writeoff_row, col_idx]
        
        if pd.notna(bus_val) and isinstance(bus_val, (int, float)):
            bus_total_calculated += float(bus_val)
        if pd.notna(ar_val) and isinstance(ar_val, (int, float)):
            ar_total_calculated += float(ar_val)
        if pd.notna(writeoff_val) and isinstance(writeoff_val, (int, float)):
            writeoff_total_calculated += float(writeoff_val)
    
    # Now check if the last column has a different total that might be correct
    last_col_idx = len(df.columns) - 1
    bus_total_existing = None
    ar_total_existing = None 
    writeoff_total_existing = None
    
    try:
        bus_val = df.iloc[bus_row, last_col_idx]
        ar_val = df.iloc[ar_row, last_col_idx]
        writeoff_val = df.iloc[writeoff_row, last_col_idx]
        
        # Check if last column value is significantly larger than individual values
        # (suggesting it might be a total rather than just another month)
        if pd.notna(bus_val):
            try:
                bus_float = float(bus_val)
                # If the last column value is much larger than the calculated total from other columns
                # and it's larger than any individual month value, it might be the real total
                if (bus_float > bus_total_calculated * 1.5 and 
                    bus_float > 0 and
                    all(bus_float > abs(float(df.iloc[bus_row, col_idx])) * 2 
                        for col_idx in numeric_cols[:-1] 
                        if pd.notna(df.iloc[bus_row, col_idx]))):
                    bus_total_existing = bus_float
            except (ValueError, TypeError):
                pass
                
        if pd.notna(ar_val):
            try:
                ar_total_existing = float(ar_val)
            except (ValueError, TypeError):
                ar_total_existing = 0.0
                
        if pd.notna(writeoff_val):
            try:
                writeoff_total_existing = float(writeoff_val)
            except (ValueError, TypeError):
                writeoff_total_existing = 0.0
    except Exception:
        pass
    
    # Use the calculated totals, but prefer existing totals if they seem more accurate
    if bus_total_existing is not None:
        bus_total = bus_total_existing
        ar_total = ar_total_existing if ar_total_existing is not None else ar_total_calculated
        writeoff_total = writeoff_total_existing if writeoff_total_existing is not None else writeoff_total_calculated
    else:
        bus_total = bus_total_calculated
        ar_total = ar_total_calculated
        writeoff_total = writeoff_total_calculated
    
    nacv_total = bus_total - ar_total - writeoff_total
    
    return {
        'bus_total': bus_total,
        'ar_total': ar_total,
        'writeoff_total': writeoff_total,
        'nacv_total': nacv_total,
        'total_losses': ar_total + writeoff_total
    }

def format_currency(v):
    return f"£{v:,.2f}"

# Flask App
app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev-secret-key")

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        excel = request.files.get("excel_file")
        contract = request.files.get("contract_file")
        chd_performance = request.form.get("chd_performance", "")

        if not excel or not contract:
            flash("Please upload both the Excel and Word files.", "error")
            return redirect(url_for("index"))

        upload_dir = os.path.join("uploads")
        os.makedirs(upload_dir, exist_ok=True)
        excel_path = os.path.join(upload_dir, "calculations.xlsx")
        contract_path = os.path.join(upload_dir, "contract.docx")
        excel.save(excel_path)
        contract.save(contract_path)

        try:
            # Calculate monthly totals
            totals = calculate_monthly_totals(excel_path)
            
            # Extract tiers and CHD rules from Word contract
            tiers = extract_tiers_from_docx(contract_path)
            chd_rules = extract_chd_rules_from_docx(contract_path)
            
            # Calculate gross incentive
            gross, breakdown = compute_gross_incentive(totals['nacv_total'], tiers)
            
            # Apply CHD adjustment if CHD performance is provided
            chd_adjustment = 0.0
            chd_breakdown = {"applied": False}
            net_incentive = gross
            
            if chd_performance.strip():
                try:
                    chd_days = float(chd_performance.strip())
                    chd_adjustment, chd_breakdown = compute_chd_adjustment(gross, chd_days, chd_rules)
                    
                    # Apply adjustment based on type (deduct or add)
                    if chd_breakdown.get("adjustment_type") == "add":
                        net_incentive = gross + chd_adjustment
                    else:  # Default to deduct
                        net_incentive = gross - chd_adjustment
                        
                except ValueError:
                    flash("Invalid CHD performance value. Please enter a number.", "error")
            
            # Build tiers view for display
            tiers_view = []
            for t in tiers:
                lo = t["lower_m"]
                hi = t["upper_m"]
                bps = t["bps"]
                if hi is None:
                    tiers_view.append(f"{lo:.2f}M+ : {bps:.1f} bps")
                else:
                    tiers_view.append(f"{lo:.2f}M – {hi:.2f}M : {bps:.1f} bps")

            return render_template(
                "result.html",
                tiers=tiers_view,
                nacv_value=totals['nacv_total'],
                gross_value=gross,
                breakdown=breakdown,
                chd_rules=chd_rules,
                chd_adjustment=chd_adjustment,
                chd_breakdown=chd_breakdown,
                net_incentive=net_incentive,
                total_bus=totals['bus_total'],
                total_losses=totals['total_losses'],
                bus_col="Bus Unadjusted wo cash - Client",
                ar_col="AR 180 Days Loss", 
                wr_col="Writeoffs - New Reserve Debits",
                format_currency=format_currency
            )
            
        except Exception as e:
            flash(f"Error processing files: {e}", "error")
            return redirect(url_for("index"))

    return render_template("index.html")

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)