"""
Enhanced Document Processing Module
Implements advanced table structure recognition, fuzzy matching, and context analysis
"""

import os
import re
import logging
from typing import Dict, List, Tuple, Optional, Any
from dataclasses import dataclass
import pandas as pd
import numpy as np
from docx import Document
from fuzzywuzzy import fuzz, process
import camelot
import tabula
import spacy
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Download required NLTK data
try:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
    nltk.download('averaged_perceptron_tagger', quiet=True)
except:
    logger.warning("Could not download NLTK data")

# Load spaCy model (fallback to smaller model if large one not available)
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    logger.warning("spaCy model not found. Run: python -m spacy download en_core_web_sm")
    nlp = None

@dataclass
class TableDetectionResult:
    """Result from table detection with confidence scoring"""
    tables: List[pd.DataFrame]
    confidence: float
    method_used: str
    metadata: Dict[str, Any]

@dataclass
class FuzzyMatchResult:
    """Result from fuzzy matching with context"""
    match: str
    score: float
    context: str
    position: Tuple[int, int]  # (paragraph_index, character_position)

class EnhancedDocumentProcessor:
    """Advanced document processor with table recognition, fuzzy matching, and context analysis"""
    
    def __init__(self, fuzzy_threshold: int = 85):
        self.fuzzy_threshold = fuzzy_threshold
        self.stop_words = set(stopwords.words('english')) if nltk else set()
        
        # Enhanced keyword patterns with variations
        self.bir_keywords = {
            'table_headers': ['nacv', 'cc-nacv', 'charge volume', 'usd mm', 'volume tier', 'tier volume'],
            'rate_headers': ['bir', 'bp', 'basis point', 'bps', 'rate', 'multiplier'],
            'variants': ['cc nacv', 'net adjusted charge volume', 'basis points', 'bp multiplier']
        }
        
        self.chd_keywords = {
            'identifiers': ['chd', 'client held day', 'held day', 'client held days'],
            'thresholds': ['days or more', 'days or greater', 'minimum days', 'threshold'],
            'adjustments': ['bps per day', 'basis points per day', 'bp per day', 'deduction', 'penalty'],
            'context_words': ['performance', 'adjustment', 'incentive', 'calculation']
        }
        
        self.excel_patterns = {
            'cv_tier': ['cv tier', 'tier', 'lo-roc', 'loroc', 'lo roc'],
            'nacv': ['bus unadjusted', 'ar 180', 'writeoff', 'reserve debit', 'net adjusted'],
            'financial_terms': ['volume', 'amount', 'loss', 'charge', 'revenue']
        }

    def extract_tables_with_structure_recognition(self, file_path: str) -> TableDetectionResult:
        """
        Extract tables using multiple methods with structure recognition
        """
        tables = []
        confidence = 0.0
        method_used = "none"
        metadata = {}
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                # Try Camelot first (better for complex tables)
                try:
                    camelot_tables = camelot.read_pdf(file_path, pages='all')
                    if len(camelot_tables) > 0:
                        tables = [table.df for table in camelot_tables]
                        confidence = self._calculate_table_confidence(tables)
                        method_used = "camelot"
                        metadata = {
                            'parsing_report': [table.parsing_report for table in camelot_tables],
                            'accuracy': [table.accuracy for table in camelot_tables]
                        }
                        logger.info(f"Camelot extracted {len(tables)} tables")
                except Exception as e:
                    logger.warning(f"Camelot failed: {e}")
                
                # Fallback to Tabula
                if not tables:
                    try:
                        tabula_tables = tabula.read_pdf(file_path, pages='all', multiple_tables=True)
                        if tabula_tables:
                            tables = tabula_tables
                            confidence = self._calculate_table_confidence(tables)
                            method_used = "tabula"
                            metadata = {'table_count': len(tables)}
                            logger.info(f"Tabula extracted {len(tables)} tables")
                    except Exception as e:
                        logger.warning(f"Tabula failed: {e}")
                        
            elif file_ext in ['.xlsx', '.xls']:
                # Enhanced Excel processing
                try:
                    xls = pd.ExcelFile(file_path)
                    sheet_tables = []
                    for sheet_name in xls.sheet_names:
                        df = pd.read_excel(file_path, sheet_name=sheet_name, header=None)
                        # Detect table boundaries
                        table_regions = self._detect_table_regions(df)
                        for region in table_regions:
                            table = self._extract_table_from_region(df, region)
                            if not table.empty:
                                sheet_tables.append(table)
                    
                    tables = sheet_tables
                    confidence = self._calculate_table_confidence(tables)
                    method_used = "pandas_enhanced"
                    metadata = {'sheets_processed': len(xls.sheet_names), 'regions_found': len(table_regions)}
                    logger.info(f"Enhanced Excel processing found {len(tables)} table regions")
                except Exception as e:
                    logger.error(f"Enhanced Excel processing failed: {e}")
                    
        except Exception as e:
            logger.error(f"Table extraction failed: {e}")
            
        return TableDetectionResult(tables, confidence, method_used, metadata)

    def _detect_table_regions(self, df: pd.DataFrame) -> List[Dict]:
        """
        Detect potential table regions in Excel sheet using data density analysis
        """
        regions = []
        
        # Find regions with high data density
        for start_row in range(0, len(df), 5):  # Check every 5 rows
            for start_col in range(0, len(df.columns), 3):  # Check every 3 columns
                # Define a potential table region
                end_row = min(start_row + 20, len(df))
                end_col = min(start_col + 10, len(df.columns))
                
                region = df.iloc[start_row:end_row, start_col:end_col]
                
                # Calculate data density (non-null cells)
                non_null_ratio = region.count().sum() / (region.shape[0] * region.shape[1])
                
                if non_null_ratio > 0.3:  # At least 30% filled
                    regions.append({
                        'start_row': start_row,
                        'end_row': end_row,
                        'start_col': start_col,
                        'end_col': end_col,
                        'density': non_null_ratio
                    })
        
        # Sort by density and remove overlapping regions
        regions.sort(key=lambda x: x['density'], reverse=True)
        return self._remove_overlapping_regions(regions)

    def _remove_overlapping_regions(self, regions: List[Dict]) -> List[Dict]:
        """Remove overlapping table regions, keeping the ones with higher density"""
        filtered = []
        
        for region in regions:
            overlaps = False
            for existing in filtered:
                if self._regions_overlap(region, existing):
                    overlaps = True
                    break
            if not overlaps:
                filtered.append(region)
                
        return filtered

    def _regions_overlap(self, region1: Dict, region2: Dict) -> bool:
        """Check if two regions overlap significantly"""
        r1_area = (region1['end_row'] - region1['start_row']) * (region1['end_col'] - region1['start_col'])
        r2_area = (region2['end_row'] - region2['start_row']) * (region2['end_col'] - region2['start_col'])
        
        # Calculate overlap
        overlap_rows = max(0, min(region1['end_row'], region2['end_row']) - max(region1['start_row'], region2['start_row']))
        overlap_cols = max(0, min(region1['end_col'], region2['end_col']) - max(region1['start_col'], region2['start_col']))
        overlap_area = overlap_rows * overlap_cols
        
        # If overlap is more than 50% of smaller region, consider them overlapping
        min_area = min(r1_area, r2_area)
        return (overlap_area / min_area) > 0.5 if min_area > 0 else False

    def _extract_table_from_region(self, df: pd.DataFrame, region: Dict) -> pd.DataFrame:
        """Extract table from detected region with header detection"""
        table = df.iloc[region['start_row']:region['end_row'], 
                        region['start_col']:region['end_col']].copy()
        
        # Try to detect headers
        header_row = self._detect_header_row(table)
        if header_row is not None:
            # Set the detected row as header
            new_header = table.iloc[header_row]
            table = table.iloc[header_row+1:].copy()
            table.columns = new_header
            
        return table.dropna(how='all').dropna(axis=1, how='all')

    def _detect_header_row(self, table: pd.DataFrame) -> Optional[int]:
        """Detect which row is most likely the header"""
        for i in range(min(5, len(table))):  # Check first 5 rows
            row = table.iloc[i]
            # Check if row contains text (potential headers)
            text_cells = sum(1 for cell in row if isinstance(cell, str) and len(str(cell).strip()) > 0)
            if text_cells >= len(row) * 0.5:  # At least 50% text cells
                return i
        return None

    def _calculate_table_confidence(self, tables: List[pd.DataFrame]) -> float:
        """Calculate confidence score based on table structure quality"""
        if not tables:
            return 0.0
        
        scores = []
        for table in tables:
            # Factors: size, data density, column variety
            size_score = min(1.0, (table.shape[0] * table.shape[1]) / 100)
            density_score = table.count().sum() / (table.shape[0] * table.shape[1])
            
            # Column type variety (numeric vs text)
            numeric_cols = sum(1 for col in table.columns if table[col].dtype in ['int64', 'float64'])
            variety_score = min(1.0, numeric_cols / max(1, len(table.columns)))
            
            table_score = (size_score + density_score + variety_score) / 3
            scores.append(table_score)
            
        return sum(scores) / len(scores)

    def fuzzy_match_keywords(self, text: str, keyword_groups: Dict[str, List[str]], 
                           context_window: int = 50) -> List[FuzzyMatchResult]:
        """
        Perform fuzzy matching with context analysis
        """
        results = []
        sentences = sent_tokenize(text) if nltk else [text]
        
        for sent_idx, sentence in enumerate(sentences):
            sentence_lower = sentence.lower()
            
            for group_name, keywords in keyword_groups.items():
                for keyword in keywords:
                    # Use fuzzy matching to find similar terms
                    words_in_sentence = word_tokenize(sentence_lower) if nltk else sentence_lower.split()
                    
                    for word_idx, word in enumerate(words_in_sentence):
                        ratio = fuzz.ratio(keyword.lower(), word)
                        partial_ratio = fuzz.partial_ratio(keyword.lower(), word)
                        
                        # Also check phrase matching for multi-word keywords
                        if len(keyword.split()) > 1:
                            phrase_start = max(0, word_idx - len(keyword.split()) + 1)
                            phrase_end = min(len(words_in_sentence), word_idx + len(keyword.split()))
                            phrase = ' '.join(words_in_sentence[phrase_start:phrase_end])
                            phrase_ratio = fuzz.ratio(keyword.lower(), phrase)
                            ratio = max(ratio, phrase_ratio)
                        
                        max_ratio = max(ratio, partial_ratio)
                        
                        if max_ratio >= self.fuzzy_threshold:
                            # Extract context around the match
                            context_start = max(0, sentence.lower().find(word) - context_window)
                            context_end = min(len(sentence), sentence.lower().find(word) + len(word) + context_window)
                            context = sentence[context_start:context_end].strip()
                            
                            results.append(FuzzyMatchResult(
                                match=word,
                                score=max_ratio,
                                context=context,
                                position=(sent_idx, sentence.lower().find(word))
                            ))
        
        # Remove duplicates and sort by score
        unique_results = {}
        for result in results:
            key = (result.match, result.position)
            if key not in unique_results or unique_results[key].score < result.score:
                unique_results[key] = result
                
        return sorted(unique_results.values(), key=lambda x: x.score, reverse=True)

    def extract_context_aware_chd_rules(self, docx_path: str) -> Dict:
        """
        Enhanced CHD rule extraction with context analysis and fuzzy matching
        """
        doc = Document(docx_path)
        full_text = ' '.join([para.text for para in doc.paragraphs])
        
        # Use fuzzy matching to find CHD-related content
        chd_matches = self.fuzzy_match_keywords(full_text, self.chd_keywords)
        
        chd_rules = {
            "threshold_days": None,
            "adjustment_bps_per_day": None,
            "description": "",
            "adjustment_type": None,
            "confidence": 0.0,
            "context_analysis": {},
            "matches_found": len(chd_matches)
        }
        
        if not chd_matches:
            return chd_rules
        
        # Analyze contexts for rule extraction
        contexts = [match.context for match in chd_matches if match.score > 90]
        
        for context in contexts:
            # Enhanced threshold detection with fuzzy patterns
            threshold_patterns = [
                r'(\d+(?:\.\d+)?)\s*days?\s+or\s+more',
                r'is\s+(\d+(?:\.\d+)?)\s*days?\s+or\s+more',
                r'(\d+(?:\.\d+)?)\s*days?\s+or\s+greater',
                r'>=?\s*(\d+(?:\.\d+)?)\s*days?',
                r'(\d+(?:\.\d+)?)\s*days?\s+or\s+above',
                r'minimum\s+(\d+(?:\.\d+)?)\s*days?',
                r'threshold\s+of\s+(\d+(?:\.\d+)?)\s*days?'
            ]
            
            for pattern in threshold_patterns:
                match = re.search(pattern, context.lower())
                if match:
                    chd_rules["threshold_days"] = float(match.group(1))
                    break
            
            # Enhanced BPS detection
            bps_patterns = [
                r'(\d+(?:\.\d+)?)\s*bps?\s+per\s+day',
                r'(\d+(?:\.\d+)?)\s*basis\s+points?\s+per\s+day',
                r'(\d+(?:\.\d+)?)\s*bp\s+per\s+day',
                r'(\d+(?:\.\d+)?)\s*basis\s+point\s+per\s+day',
                r'(\d+(?:\.\d+)?)\s*bps?\s+for\s+each\s+day'
            ]
            
            for pattern in bps_patterns:
                match = re.search(pattern, context.lower())
                if match:
                    chd_rules["adjustment_bps_per_day"] = float(match.group(1))
                    break
            
            # Enhanced sentiment analysis for adjustment type
            deduct_indicators = ['deduct', 'subtract', 'reduce', 'penalty', 'loss', 'negative', 'decrease']
            add_indicators = ['add', 'bonus', 'increase', 'positive', 'benefit', 'gain']
            
            context_lower = context.lower()
            deduct_score = sum(1 for word in deduct_indicators if word in context_lower)
            add_score = sum(1 for word in add_indicators if word in context_lower)
            
            if deduct_score > add_score:
                chd_rules["adjustment_type"] = "deduct"
            elif add_score > deduct_score:
                chd_rules["adjustment_type"] = "add"
            
            # Store the most informative context
            if len(context) > len(chd_rules["description"]):
                chd_rules["description"] = context
        
        # Calculate confidence based on completeness and match scores
        completeness = sum([
            1 if chd_rules["threshold_days"] is not None else 0,
            1 if chd_rules["adjustment_bps_per_day"] is not None else 0,
            1 if chd_rules["adjustment_type"] is not None else 0
        ]) / 3
        
        avg_match_score = sum(match.score for match in chd_matches) / len(chd_matches) if chd_matches else 0
        chd_rules["confidence"] = (completeness + avg_match_score / 100) / 2
        
        # Add context analysis
        chd_rules["context_analysis"] = {
            "high_confidence_matches": len([m for m in chd_matches if m.score > 95]),
            "medium_confidence_matches": len([m for m in chd_matches if 85 <= m.score <= 95]),
            "contexts_analyzed": len(contexts),
            "avg_match_score": avg_match_score
        }
        
        return chd_rules

    def extract_enhanced_bir_table(self, docx_path: str) -> List[Dict]:
        """
        Enhanced BIR table extraction with structure recognition and fuzzy matching
        """
        # First try enhanced table extraction
        table_result = self.extract_tables_with_structure_recognition(docx_path)
        
        if table_result.tables:
            return self._process_bir_tables_from_dataframes(table_result.tables)
        
        # Fallback to document-based extraction with fuzzy matching
        return self._extract_bir_from_document_with_fuzzy_matching(docx_path)

    def _process_bir_tables_from_dataframes(self, tables: List[pd.DataFrame]) -> List[Dict]:
        """Process extracted tables to find BIR information"""
        bir_data = []
        
        for table in tables:
            # Use fuzzy matching to identify BIR tables
            header_text = ' '.join([str(col) for col in table.columns]).lower()
            
            # Check if this table contains BIR information
            volume_match = process.extractOne(
                header_text, 
                self.bir_keywords['table_headers'],
                scorer=fuzz.partial_ratio
            )
            
            rate_match = process.extractOne(
                header_text,
                self.bir_keywords['rate_headers'],
                scorer=fuzz.partial_ratio
            )
            
            if (volume_match and volume_match[1] > 70) and (rate_match and rate_match[1] > 70):
                # This looks like a BIR table
                bir_entries = self._extract_bir_from_table(table)
                bir_data.extend(bir_entries)
        
        return self._deduplicate_and_sort_bir_data(bir_data)

    def _extract_bir_from_table(self, table: pd.DataFrame) -> List[Dict]:
        """Extract BIR data from a structured table"""
        bir_entries = []
        
        for _, row in table.iterrows():
            row_values = [str(val) for val in row.values if pd.notna(val)]
            
            if len(row_values) < 2:
                continue
            
            # Try to extract range and BPS from the row
            range_text = row_values[0]
            bps_text = row_values[-1]  # Assume BPS is in the last column
            
            # Parse range
            lo, hi = self._parse_range_with_fuzzy_logic(range_text)
            
            # Parse BPS
            bps = self._parse_bps_with_fuzzy_logic(bps_text)
            
            if lo is not None and bps is not None:
                bir_entries.append({
                    "lower_m": lo,
                    "upper_m": hi,
                    "bps": bps
                })
        
        return bir_entries

    def _parse_range_with_fuzzy_logic(self, text: str) -> Tuple[Optional[float], Optional[float]]:
        """Enhanced range parsing with fuzzy logic"""
        if not text or pd.isna(text):
            return None, None
        
        text_clean = str(text).strip()
        # Remove common currency symbols and units
        text_clean = re.sub(r'[$£€¥]', '', text_clean)
        text_clean = re.sub(r'MM|million|M', '', text_clean, flags=re.IGNORECASE)
        
        # Normalize different dash types
        text_clean = text_clean.replace("–", "-").replace("—", "-").replace("−", "-")
        
        # Check for "+" indicating open-ended range
        plus_end = "+" in text_clean
        
        # Extract numbers
        numbers = re.findall(r'(\d+(?:\.\d+)?)', text_clean.replace(",", ""))
        
        if plus_end and numbers:
            return float(numbers[0]), None
        elif len(numbers) >= 2:
            return float(numbers[0]), float(numbers[1])
        elif len(numbers) == 1:
            return float(numbers[0]), None
        
        return None, None

    def _parse_bps_with_fuzzy_logic(self, text: str) -> Optional[float]:
        """Enhanced BPS parsing with fuzzy logic"""
        if not text or pd.isna(text):
            return None
        
        text_clean = str(text).strip().lower()
        
        # Remove common BPS indicators
        text_clean = re.sub(r'\b(bps?|basis\s+points?|bp)\b', '', text_clean)
        
        # Extract number
        numbers = re.findall(r'(\d+(?:\.\d+)?)', text_clean.replace(",", ""))
        
        if numbers:
            return float(numbers[0])
        
        return None

    def _extract_bir_from_document_with_fuzzy_matching(self, docx_path: str) -> List[Dict]:
        """Fallback method using fuzzy matching on document paragraphs"""
        doc = Document(docx_path)
        full_text = ' '.join([para.text for para in doc.paragraphs])
        
        # Find BIR-related sections
        bir_matches = self.fuzzy_match_keywords(full_text, self.bir_keywords)
        
        bir_data = []
        
        # Process high-confidence matches
        for match in bir_matches:
            if match.score > 85:
                # Try to extract structured data from the context
                context = match.context
                entries = self._extract_bir_entries_from_text(context)
                bir_data.extend(entries)
        
        return self._deduplicate_and_sort_bir_data(bir_data)

    def _extract_bir_entries_from_text(self, text: str) -> List[Dict]:
        """Extract BIR entries from unstructured text"""
        entries = []
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Look for patterns like "0-9.9MM: 50bps"
            pattern = r'([\d\.\-\+\s]+(?:MM|million|M)?)\s*:?\s*([\d\.]+)\s*bps?'
            matches = re.finditer(pattern, line, re.IGNORECASE)
            
            for match in matches:
                range_text = match.group(1).strip()
                bps_text = match.group(2).strip()
                
                lo, hi = self._parse_range_with_fuzzy_logic(range_text)
                bps = self._parse_bps_with_fuzzy_logic(bps_text)
                
                if lo is not None and bps is not None:
                    entries.append({
                        "lower_m": lo,
                        "upper_m": hi,
                        "bps": bps
                    })
        
        return entries

    def _deduplicate_and_sort_bir_data(self, bir_data: List[Dict]) -> List[Dict]:
        """Remove duplicates and sort BIR data"""
        if not bir_data:
            return []
        
        # Deduplicate
        unique_data = {}
        for entry in bir_data:
            key = (
                round(entry["lower_m"], 4),
                None if entry["upper_m"] is None else round(entry["upper_m"], 4),
                round(entry["bps"], 4)
            )
            unique_data[key] = entry
        
        # Sort by lower bound
        sorted_data = list(unique_data.values())
        sorted_data.sort(key=lambda x: (x["lower_m"], x["upper_m"] if x["upper_m"] is not None else float('inf')))
        
        return sorted_data

    def enhanced_excel_detection(self, excel_path: str) -> Dict:
        """
        Enhanced Excel file analysis with fuzzy matching and context analysis
        """
        try:
            df = pd.read_excel(excel_path, engine="openpyxl", header=None)
            first_column = df.iloc[:, 0].astype(str).fillna('')
            
            detection_results = {
                'detected_type': 'unknown',
                'confidence': 0.0,
                'reasoning': '',
                'indicators': [],
                'fuzzy_matches': [],
                'context_analysis': {}
            }
            
            # Combine all text for context analysis
            full_text = ' '.join(first_column.tolist())
            
            # Fuzzy match against known patterns
            cv_matches = self.fuzzy_match_keywords(full_text, {'cv_patterns': self.excel_patterns['cv_tier']})
            nacv_matches = self.fuzzy_match_keywords(full_text, {'nacv_patterns': self.excel_patterns['nacv']})
            
            cv_score = sum(match.score for match in cv_matches) / len(cv_matches) if cv_matches else 0
            nacv_score = sum(match.score for match in nacv_matches) / len(nacv_matches) if nacv_matches else 0
            
            # Determine detection type based on scores and match counts
            if cv_score > nacv_score and cv_matches:
                detection_results['detected_type'] = 'cv_tier'
                detection_results['confidence'] = min(cv_score / 100, 1.0)
                detection_results['fuzzy_matches'] = cv_matches
                detection_results['reasoning'] = f"CV Tier patterns detected with {len(cv_matches)} matches (avg score: {cv_score:.1f})"
            elif nacv_score > cv_score and nacv_matches:
                detection_results['detected_type'] = 'nacv_based'
                detection_results['confidence'] = min(nacv_score / 100, 1.0)
                detection_results['fuzzy_matches'] = nacv_matches
                detection_results['reasoning'] = f"NACV patterns detected with {len(nacv_matches)} matches (avg score: {nacv_score:.1f})"
            
            # Context analysis
            detection_results['context_analysis'] = {
                'total_cv_matches': len(cv_matches),
                'total_nacv_matches': len(nacv_matches),
                'cv_avg_score': cv_score,
                'nacv_avg_score': nacv_score,
                'high_confidence_matches': len([m for m in cv_matches + nacv_matches if m.score > 95])
            }
            
            return detection_results
            
        except Exception as e:
            logger.error(f"Enhanced Excel detection failed: {e}")
            return {
                'detected_type': 'unknown',
                'confidence': 0.0,
                'reasoning': f'Detection failed: {str(e)}',
                'indicators': [],
                'fuzzy_matches': [],
                'context_analysis': {}
            }