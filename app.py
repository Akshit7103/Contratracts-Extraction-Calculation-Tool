import os
import io
import re
from typing import Dict, List, Tuple, Optional

import pandas as pd
from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from werkzeug.utils import secure_filename
from docx import Document

ALLOWED_EXCEL = {"xls", "xlsx", "xlsm", "xlsb"}
ALLOWED_WORD = {"docx"}

def is_excel(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXCEL

def is_word(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_WORD

def format_currency(v: Optional[float]) -> str:
    """Format currency for display - supports both USD and GBP"""
    if v is None:
        return "—"
    # For NACV calculations, use GBP symbol; for CV Tier, use USD
    return "£{:,.2f}".format(v) if hasattr(format_currency, '_use_gbp') and format_currency._use_gbp else "${:,.2f}".format(v)

# -------------------------
# NACV-Based Functions (from AMEX version)
# -------------------------

def extract_chd_rules_from_docx(docx_path: str):
    """Extract CHD (Client Held Day) rules from Word contract"""
    doc = Document(docx_path)
    chd_rules = {
        "threshold_days": None,
        "adjustment_bps_per_day": None,
        "description": "",
        "adjustment_type": None  # "deduct" or "add"
    }
    
    # Search for CHD rules in paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        text_lower = text.lower()
        
        # Look for CHD-related content
        if 'chd' in text_lower and any(keyword in text_lower for keyword in ['days', 'bps', 'adjustment', 'performance']):
            
            # Look for threshold patterns (e.g., "14 days or more", "is 14 days or more")
            threshold_patterns = [
                r'(\d+(?:\.\d+)?)\s*days?\s+or\s+more',
                r'is\s+(\d+(?:\.\d+)?)\s*days?\s+or\s+more',
                r'(\d+(?:\.\d+)?)\s*days?\s+or\s+greater',
                r'>=?\s*(\d+(?:\.\d+)?)\s*days?',
                r'(\d+(?:\.\d+)?)\s*days?\s+or\s+above'
            ]
            
            for pattern in threshold_patterns:
                match = re.search(pattern, text_lower)
                if match:
                    chd_rules["threshold_days"] = float(match.group(1))
                    break
            
            # Look for bps adjustment patterns (e.g., "2 bps per day", "2.5 bps per day")
            bps_patterns = [
                r'(\d+(?:\.\d+)?)\s*bps?\s+per\s+day',
                r'(\d+(?:\.\d+)?)\s*basis\s+points?\s+per\s+day',
                r'(\d+(?:\.\d+)?)\s*bp\s+per\s+day'
            ]
            
            for pattern in bps_patterns:
                match = re.search(pattern, text_lower)
                if match:
                    chd_rules["adjustment_bps_per_day"] = float(match.group(1))
                    break
            
            # Determine if it's a deduction or addition
            if any(word in text_lower for word in ['deduct', 'subtract', 'reduce', 'penalty']):
                chd_rules["adjustment_type"] = "deduct"
            elif any(word in text_lower for word in ['add', 'bonus', 'increase']):
                chd_rules["adjustment_type"] = "add"
            elif 'will be' in text_lower and any(word in text_lower for word in ['deduct', 'subtract']):
                chd_rules["adjustment_type"] = "deduct"
                
            # Store the description (keep the most detailed one)
            if len(text) > len(chd_rules["description"]):
                chd_rules["description"] = text
    
    # If we found both threshold and bps but no adjustment type, assume deduction (most common)
    if chd_rules["threshold_days"] is not None and chd_rules["adjustment_bps_per_day"] is not None and chd_rules["adjustment_type"] is None:
        chd_rules["adjustment_type"] = "deduct"
    
    return chd_rules

def extract_tiers_from_docx(docx_path: str):
    """Extract tier rules from Word contract for NACV calculations"""
    doc = Document(docx_path)
    tiers = []
    
    def parse_range_cell(cell_text):
        t = " ".join(cell_text.split())
        t = t.replace("–", "-").replace("—", "-").replace("−", "-")
        if "+" in t:
            lower = re.findall(r"[-+]?\d*\.?\d+", t)
            if lower:
                return float(lower[0]), None
        parts = re.findall(r"[-+]?\d*\.?\d+", t)
        if len(parts) >= 2:
            return float(parts[0]), float(parts[1])
        return None, None

    def parse_bps_cell(cell_text):
        nums = re.findall(r"[-+]?\d*\.?\d+", cell_text)
        if nums:
            return float(nums[0])
        return None

    for tbl in doc.tables:
        header_text = " ".join([cell.text for cell in tbl.rows[0].cells]).lower()
        if ("nacv" in header_text and "bp" in header_text) or ("bp multiplier" in header_text):
            for i, row in enumerate(tbl.rows[1:], start=1):
                cells = [c.text.strip() for c in row.cells]
                if len(cells) < 2:
                    continue
                rng_text = cells[0]
                bps_text = cells[-1]
                lo, hi = parse_range_cell(rng_text)
                bps = parse_bps_cell(bps_text)
                if lo is not None and bps is not None:
                    tiers.append({
                        "lower_m": lo,
                        "upper_m": hi,
                        "bps": bps
                    })
    
    # Remove duplicates and sort
    dedup = []
    seen = set()
    for t in tiers:
        key = (round(t["lower_m"], 2), None if t["upper_m"] is None else round(t["upper_m"], 2), round(t["bps"], 2))
        if key not in seen:
            seen.add(key)
            dedup.append(t)
    
    dedup.sort(key=lambda x: (x["lower_m"], x["upper_m"] if x["upper_m"] is not None else 10**9))
    return dedup

def compute_chd_adjustment(gross_incentive_gbp: float, chd_performance_days: float, chd_rules):
    """Compute CHD adjustment to gross incentive"""
    if not chd_rules or chd_rules["threshold_days"] is None or chd_rules["adjustment_bps_per_day"] is None:
        return 0.0, {"applied": False, "reason": "No CHD rules found"}
    
    if chd_performance_days < chd_rules["threshold_days"]:
        return 0.0, {
            "applied": False, 
            "reason": f"CHD performance ({chd_performance_days} days) below threshold ({chd_rules['threshold_days']} days)"
        }
    
    # Calculate adjustment: bps per day based on contract
    adjustment_bps = chd_performance_days * chd_rules["adjustment_bps_per_day"]
    adjustment_rate = adjustment_bps / 10000.0
    adjustment_amount = gross_incentive_gbp * adjustment_rate
    
    # Determine the adjustment type
    adjustment_type = chd_rules.get("adjustment_type", "deduct")
    
    return adjustment_amount, {
        "applied": True,
        "chd_performance_days": chd_performance_days,
        "threshold_days": chd_rules["threshold_days"],
        "adjustment_bps_per_day": chd_rules["adjustment_bps_per_day"],
        "adjustment_type": adjustment_type,
        "total_adjustment_bps": adjustment_bps,
        "adjustment_amount_gbp": adjustment_amount
    }

def compute_gross_incentive(nacv_value_gbp: float, tiers):
    """Compute gross incentive using tier rules"""
    if not tiers:
        raise ValueError("No tiers found in the contract.")

    # Find top tier (with upper_m = None)
    top_idx = None
    for i, t in enumerate(tiers):
        if t["upper_m"] is None:
            top_idx = i
            break
    if top_idx is None:
        top_idx = len(tiers) - 1

    top = tiers[top_idx]
    if top_idx == 0:
        raise ValueError("Tier list malformed (no lower tiers before top).")

    second = tiers[top_idx - 1]
    second_upper_m = second["upper_m"] if second["upper_m"] is not None else second["lower_m"]

    nacv_m = nacv_value_gbp / 1_000_000.0

    if nacv_m < top["lower_m"]:
        # Single bracket calculation
        chosen = None
        for t in tiers:
            lo = t["lower_m"]
            hi = t["upper_m"]
            if hi is None:
                continue
            if nacv_m >= lo and nacv_m <= hi:
                chosen = t
                break
        if chosen is None:
            chosen = tiers[0]
        
        bps = chosen["bps"]
        rate = bps / 10000.0
        gross = nacv_value_gbp * rate
        breakdown = {
            "mode": "single-bracket",
            "bracket_bps": bps,
            "applied_amount_gbp": nacv_value_gbp,
            "calculation_steps": [
                f"NACV Value: {format_currency_gbp(nacv_value_gbp)}",
                f"Applied Rate: {bps:.1f} bps ({rate:.4%})",
                f"Calculation: {format_currency_gbp(nacv_value_gbp)} × {rate:.4%} = {format_currency_gbp(gross)}"
            ],
            "explanation": f"Since your NACV of {format_currency_gbp(nacv_value_gbp)} falls within the {chosen['lower_m']:.2f}M - {chosen['upper_m']:.2f}M tier, we apply a single rate of {bps:.1f} bps to the entire amount."
        }
        return gross, breakdown
    else:
        # Two-tier calculation
        second_upper_gbp = second_upper_m * 1_000_000.0
        top_bps = top["bps"]
        second_bps = second["bps"]
        top_rate = top_bps / 10000.0
        second_rate = second_bps / 10000.0

        base_amount = min(nacv_value_gbp, second_upper_gbp)
        remainder = max(0.0, nacv_value_gbp - second_upper_gbp)

        base_incentive = base_amount * second_rate
        remainder_incentive = remainder * top_rate
        gross = base_incentive + remainder_incentive
        
        breakdown = {
            "mode": "two-tier",
            "second_upper_m": second_upper_m,
            "second_bps": second_bps,
            "top_bps": top_bps,
            "base_amount_gbp": base_amount,
            "remainder_amount_gbp": remainder,
            "base_incentive": base_incentive,
            "remainder_incentive": remainder_incentive,
            "calculation_steps": [
                f"NACV Value: {format_currency_gbp(nacv_value_gbp)}",
                f"Tier 1 (up to {second_upper_m:.2f}M): {format_currency_gbp(base_amount)} × {second_bps:.1f} bps ({second_rate:.4%}) = {format_currency_gbp(base_incentive)}",
                f"Tier 2 (above {second_upper_m:.2f}M): {format_currency_gbp(remainder)} × {top_bps:.1f} bps ({top_rate:.4%}) = {format_currency_gbp(remainder_incentive)}",
                f"Total Gross Incentive: {format_currency_gbp(base_incentive)} + {format_currency_gbp(remainder_incentive)} = {format_currency_gbp(gross)}"
            ],
            "explanation": f"Your NACV of {format_currency_gbp(nacv_value_gbp)} exceeds {second_upper_m:.2f}M, so we use a two-tier calculation: the first {format_currency_gbp(second_upper_gbp)} at {second_bps:.1f} bps, and the remaining {format_currency_gbp(remainder)} at {top_bps:.1f} bps."
        }
        return gross, breakdown

def calculate_monthly_totals(excel_path: str):
    """
    Simple function to calculate totals from monthly data.
    Expects horizontal format: rows = categories, columns = months
    """
    df = pd.read_excel(excel_path, engine="openpyxl")
    
    # Find the rows for each category
    bus_row = None
    ar_row = None
    writeoff_row = None
    
    for i, row_label in enumerate(df.iloc[:, 0]):
        if pd.isna(row_label):
            continue
        label_lower = str(row_label).lower()
        if "bus unadjusted" in label_lower:
            bus_row = i
        elif "ar 180" in label_lower:
            ar_row = i
        elif "writeoff" in label_lower and "reserve" in label_lower:
            writeoff_row = i
    
    if bus_row is None or ar_row is None or writeoff_row is None:
        raise ValueError("Could not find required data rows in Excel file")
    
    # First, calculate totals from individual month columns
    numeric_cols = []
    for col_idx in range(1, len(df.columns)):  # Skip first column (labels)
        if pd.api.types.is_numeric_dtype(df.iloc[:, col_idx]):
            numeric_cols.append(col_idx)
    
    bus_total_calculated = 0.0
    ar_total_calculated = 0.0
    writeoff_total_calculated = 0.0
    
    for col_idx in numeric_cols:
        bus_val = df.iloc[bus_row, col_idx]
        ar_val = df.iloc[ar_row, col_idx]
        writeoff_val = df.iloc[writeoff_row, col_idx]
        
        if pd.notna(bus_val) and isinstance(bus_val, (int, float)):
            bus_total_calculated += float(bus_val)
        if pd.notna(ar_val) and isinstance(ar_val, (int, float)):
            ar_total_calculated += float(ar_val)
        if pd.notna(writeoff_val) and isinstance(writeoff_val, (int, float)):
            writeoff_total_calculated += float(writeoff_val)
    
    # Now check if the last column has a different total that might be correct
    last_col_idx = len(df.columns) - 1
    bus_total_existing = None
    ar_total_existing = None 
    writeoff_total_existing = None
    
    try:
        bus_val = df.iloc[bus_row, last_col_idx]
        ar_val = df.iloc[ar_row, last_col_idx]
        writeoff_val = df.iloc[writeoff_row, last_col_idx]
        
        # Check if last column value is significantly larger than individual values
        # (suggesting it might be a total rather than just another month)
        if pd.notna(bus_val):
            try:
                bus_float = float(bus_val)
                # If the last column value is much larger than the calculated total from other columns
                # and it's larger than any individual month value, it might be the real total
                if (bus_float > bus_total_calculated * 1.5 and 
                    bus_float > 0 and
                    all(bus_float > abs(float(df.iloc[bus_row, col_idx])) * 2 
                        for col_idx in numeric_cols[:-1] 
                        if pd.notna(df.iloc[bus_row, col_idx]))):
                    bus_total_existing = bus_float
            except (ValueError, TypeError):
                pass
                
        if pd.notna(ar_val):
            try:
                ar_total_existing = float(ar_val)
            except (ValueError, TypeError):
                ar_total_existing = 0.0
                
        if pd.notna(writeoff_val):
            try:
                writeoff_total_existing = float(writeoff_val)
            except (ValueError, TypeError):
                writeoff_total_existing = 0.0
    except Exception:
        pass
    
    # Use the calculated totals, but prefer existing totals if they seem more accurate
    if bus_total_existing is not None:
        bus_total = bus_total_existing
        ar_total = ar_total_existing if ar_total_existing is not None else ar_total_calculated
        writeoff_total = writeoff_total_existing if writeoff_total_existing is not None else writeoff_total_calculated
    else:
        bus_total = bus_total_calculated
        ar_total = ar_total_calculated
        writeoff_total = writeoff_total_calculated
    
    nacv_total = bus_total - ar_total - writeoff_total
    
    return {
        'bus_total': bus_total,
        'ar_total': ar_total,
        'writeoff_total': writeoff_total,
        'nacv_total': nacv_total,
        'total_losses': ar_total + writeoff_total
    }

def format_currency_gbp(v):
    """Helper function for GBP formatting"""
    return f"£{v:,.2f}"

# -------------------------
# Auto-Detection Functions
# -------------------------

def detect_calculation_type(excel_path: str) -> Dict:
    """
    Auto-detect whether the Excel file is CV Tier-based or NACV-based
    Priority: Look for 'tier' keyword first for tier-based detection
    Returns: {
        'detected_type': 'cv_tier' | 'nacv_based' | 'unknown',
        'confidence': 'high' | 'medium' | 'low',
        'reasoning': 'explanation of detection logic',
        'indicators': ['list', 'of', 'found', 'indicators']
    }
    """
    try:
        df = pd.read_excel(excel_path, engine="openpyxl", header=None)
        first_column = df.iloc[:, 0].astype(str).str.lower().fillna('')
        
        # Step 1: Check for "tier" keyword first (highest priority for CV Tier detection)
        tier_matches = []
        tier_rows = []
        for idx, cell in enumerate(first_column):
            if "tier" in cell:
                # Extract the actual tier text
                tier_text = df.iloc[idx, 0] if pd.notna(df.iloc[idx, 0]) else str(cell)
                tier_matches.append(str(tier_text).strip())
                tier_rows.append(f"row {idx + 1}")
        
        # If we found "tier" keywords, this is very likely CV Tier-based
        if len(tier_matches) >= 1:
            confidence = "high" if len(tier_matches) >= 2 else "medium"
            reasoning = f"Found {len(tier_matches)} 'tier' references: {', '.join(tier_matches[:3])} in {', '.join(tier_rows[:3])}"
            return {
                'detected_type': 'cv_tier',
                'confidence': confidence,
                'reasoning': reasoning,
                'indicators': tier_matches
            }
        
        # Step 2: Check for specific NACV indicators (higher priority than generic CV terms)
        nacv_indicators = {
            "bus unadjusted": "Bus Unadjusted",
            "ar 180": "AR 180 Days", 
            "writeoff": "Writeoffs",
            "reserve debit": "Reserve Debits",
            "net adjusted": "Net Adjusted"
        }
        
        nacv_matches = []
        nacv_rows = []
        for keyword, display_name in nacv_indicators.items():
            for idx, cell in enumerate(first_column):
                if keyword in cell:
                    nacv_matches.append(display_name)
                    nacv_rows.append(f"row {idx + 1}")
                    break
        
        # Strong NACV indicators found
        if len(nacv_matches) >= 2:
            confidence = "high" if len(nacv_matches) >= 3 else "medium"
            reasoning = f"Found {len(nacv_matches)} NACV indicators: {', '.join(nacv_matches)} in {', '.join(nacv_rows[:3])}"
            return {
                'detected_type': 'nacv_based',
                'confidence': confidence,
                'reasoning': reasoning,
                'indicators': nacv_matches
            }
        
        # Step 3: Check for other CV Tier indicators (lower priority)
        cv_tier_indicators = {
            "cv tier": "CV Tier",
            "lo-roc": "Lo-ROC",
            "loroc": "LoROC", 
            "lo roc": "Lo ROC"
        }
        
        cv_tier_matches = []
        cv_tier_rows = []
        for keyword, display_name in cv_tier_indicators.items():
            for idx, cell in enumerate(first_column):
                if keyword in cell:
                    cv_tier_matches.append(display_name)
                    cv_tier_rows.append(f"row {idx + 1}")
                    break
        
        if len(cv_tier_matches) >= 1:
            confidence = "high" if len(cv_tier_matches) >= 2 else "medium"
            reasoning = f"Found {len(cv_tier_matches)} CV Tier indicators: {', '.join(cv_tier_matches)} in {', '.join(cv_tier_rows[:3])}"
            return {
                'detected_type': 'cv_tier',
                'confidence': confidence,
                'reasoning': reasoning,
                'indicators': cv_tier_matches
            }
        
        # Step 4: Single NACV indicator (low confidence)
        if len(nacv_matches) == 1:
            reasoning = f"Found 1 NACV indicator: {nacv_matches[0]} (low confidence - please verify)"
            return {
                'detected_type': 'nacv_based',
                'confidence': 'low',
                'reasoning': reasoning,
                'indicators': nacv_matches
            }
        
        # Step 5: No clear indicators found
        reasoning = "No clear CV Tier or NACV indicators found in first column"
        return {
            'detected_type': 'unknown',
            'confidence': 'low',
            'reasoning': reasoning,
            'indicators': []
        }
            
    except Exception as e:
        return {
            'detected_type': 'unknown',
            'confidence': 'low',
            'reasoning': f"Error analyzing file: {str(e)}",
            'indicators': []
        }

def validate_detection_with_contract(contract_path: str, detected_type: str) -> Dict:
    """
    Validate auto-detection by analyzing the contract file
    Returns confidence boost information
    """
    try:
        doc = Document(contract_path)
        full_text = " ".join([para.text.lower() for para in doc.paragraphs])
        
        # Look for contract indicators
        chd_indicators = ["chd", "client held day", "held day"]
        bir_indicators = ["bir", "basis point", "bp multiplier"]
        nacv_table_indicators = ["nacv", "bp multiplier", "tier"]
        
        chd_found = any(indicator in full_text for indicator in chd_indicators)
        bir_found = any(indicator in full_text for indicator in bir_indicators)
        nacv_table_found = any(indicator in full_text for indicator in nacv_table_indicators)
        
        validation_result = {
            'contract_supports_detection': False,
            'contract_reasoning': '',
            'confidence_boost': 0
        }
        
        if detected_type == 'nacv_based':
            if chd_found or (nacv_table_found and not bir_found):
                validation_result['contract_supports_detection'] = True
                validation_result['contract_reasoning'] = "Contract contains CHD rules or NACV tables"
                validation_result['confidence_boost'] = 1
            elif bir_found:
                validation_result['contract_reasoning'] = "Warning: Contract contains BIR tables (more typical of CV Tier)"
                validation_result['confidence_boost'] = -1
                
        elif detected_type == 'cv_tier':
            if bir_found and not chd_found:
                validation_result['contract_supports_detection'] = True
                validation_result['contract_reasoning'] = "Contract contains BIR basis point tables"
                validation_result['confidence_boost'] = 1
            elif chd_found:
                validation_result['contract_reasoning'] = "Warning: Contract contains CHD rules (more typical of NACV)"
                validation_result['confidence_boost'] = -1
                
        return validation_result
        
    except Exception as e:
        return {
            'contract_supports_detection': False,
            'contract_reasoning': f"Could not analyze contract: {str(e)}",
            'confidence_boost': 0
        }

# -------------------------
# CV Tier-Based Functions (Original)
# -------------------------
def _num(text: str) -> Optional[float]:
    if text is None:
        return None
    m = re.findall(r"[-+]?\d*\.?\d+", str(text).replace(",", ""))
    if not m:
        return None
    try:
        return float(m[0])
    except:
        return None

def extract_bir_table_from_contract(docx_path: str) -> List[Dict]:
    """
    Extracts a BIR (basis points) table from the contract .docx.
    Looks for a table whose header mentions NACV/CC-NACV/Charge Volume and bp/BIR.
    Returns list of dicts: [{'lower_m': float, 'upper_m': Optional[float], 'bps': float}, ...]
    Units: 'lower_m' and 'upper_m' are in MILLIONS; 'bps' is basis points (1bp = 0.01%).
    """
    doc = Document(docx_path)
    tiers = []

    def parse_range_cell(cell_text: str) -> Tuple[Optional[float], Optional[float]]:
        t = " ".join((cell_text or "").split())
        t = t.replace("–", "-").replace("—", "-").replace("−", "-")
        # Examples: "0 - $9.9MM", "$10.0MM - $24.9MM", "$75.0MM+"
        # Strip currency/MM/words, keep numbers
        # Identify a trailing "+" meaning open-ended top tier
        plus = "+" in t
        nums = re.findall(r"[-+]?\d*\.?\d+", t.replace(",", ""))
        if plus and nums:
            return float(nums[0]), None
        if len(nums) >= 2:
            return float(nums[0]), float(nums[1])
        if len(nums) == 1 and not plus:
            # single number w/o plus is ambiguous; treat as lower bound w/None
            return float(nums[0]), None
        return None, None

    def parse_bps_cell(cell_text: str) -> Optional[float]:
        # Accept "50.0 bp", "50 bp", "50.0 basis points", "50.0 bps"
        nums = re.findall(r"[-+]?\d*\.?\d+", (cell_text or "").replace(",", ""))
        if nums:
            try:
                return float(nums[0])
            except:
                return None
        return None

    for tbl in doc.tables:
        # construct header text
        header_text = " ".join([cell.text for cell in tbl.rows[0].cells]).lower()
        if any(k in header_text for k in ["nacv", "cc-nacv", "charge volume", "usd mm"]) and any(k in header_text for k in ["bir", "bp", "basis point"]):
            # parse body rows
            for row in tbl.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) < 2:
                    continue
                lo, hi = parse_range_cell(cells[0])
                bps = parse_bps_cell(cells[-1])
                if lo is None or bps is None:
                    continue
                tiers.append({"lower_m": lo, "upper_m": hi, "bps": bps})

    # de-duplicate and sort
    uniq = {}
    for t in tiers:
        key = (round(t["lower_m"], 4), None if t["upper_m"] is None else round(t["upper_m"], 4), round(t["bps"], 4))
        uniq[key] = t
    tiers = list(uniq.values())
    tiers.sort(key=lambda x: (x["lower_m"], x["upper_m"] if x["upper_m"] is not None else 10**9))
    return tiers


# -------------------------
# Helpers: CHD benchmark parsing
# -------------------------
def extract_chd_benchmark_from_contract(docx_path: str) -> Optional[float]:
    """
    Attempts to extract the 'CHD Benchmark' numeric value from the contract .docx.
    Searches paragraphs and table cells for the phrase 'CHD Benchmark' and returns the first number found.
    Nothing is hardcoded; uses regex to pull the numeric value.
    """
    doc = Document(docx_path)
    import re

    def find_number_near(text: str) -> Optional[float]:
        if not text:
            return None
        if "chd benchmark" in text.lower():
            nums = re.findall(r"[-+]?\d*\.?\d+", text.replace(",", ""))
            for n in nums:
                try:
                    return float(n)
                except:
                    continue
        return None

    # Search paragraphs
    for p in doc.paragraphs:
        v = find_number_near(p.text)
        if v is not None:
            return v

    # Search tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                v = find_number_near(cell.text)
                if v is not None:
                    return v

    return None
# -------------------------
# Helpers: Excel parsing
# -------------------------
def read_excel_flex(excel_path: str) -> pd.DataFrame:
    """
    Read the first sheet of an Excel file as raw data (no header inference).
    This makes the parser robust to metadata rows and merged headers.
    """
    xls = pd.ExcelFile(excel_path)
    sheet = xls.sheet_names[0]
    return pd.read_excel(excel_path, sheet_name=sheet, header=None, engine="openpyxl")

def compute_cv_tier_totals(df: pd.DataFrame) -> Dict[str, float]:
    """
    Scan the first column for rows that look like CV Tier or Lo-ROC categories.
    Sum across all numeric cells in the row (ignoring blanks/text).
    Returns mapping: {label -> total_amount}
    """
    label_col = df.iloc[:, 0].astype(str)
    categories = {}
    for idx, label in label_col.items():
        lbl = (label or "").strip()
        low = lbl.lower()
        if not lbl or lbl == "nan":
            continue
        if ("cv" in low and "tier" in low) or ("lo-roc" in low) or ("loroc" in low) or ("lo roc" in low):
            # sum numeric values across the row (skip first column)
            row = df.iloc[idx, 1:]
            vals = pd.to_numeric(row, errors="coerce")
            total = float(vals.fillna(0).sum())
            # keep original label (normalized spacing)
            lbl_norm = " ".join(lbl.split())
            categories[lbl_norm] = categories.get(lbl_norm, 0.0) + total
    return categories

# -------------------------
# Helpers: Calculation
# -------------------------
def find_bir_for_amount(amount_usd: float, bir_tiers: List[Dict]) -> Optional[float]:
    """
    Given an amount in USD, pick the matching BIR (bps) using the NACV (USD MM) table.
    """
    amount_m = amount_usd / 1_000_000.0
    chosen = None
    for t in bir_tiers:
        lo = t["lower_m"]
        hi = t["upper_m"]
        if hi is None:
            # top open-ended tier
            if amount_m >= lo:
                chosen = t
        else:
            if amount_m >= lo and amount_m <= hi:
                chosen = t
    return None if chosen is None else chosen["bps"]

def calc_per_tier_adjustments(category_totals: Dict[str, float], bir_tiers: List[Dict]) -> List[Dict]:
    """
    For each CV Tier/Lo-ROC total, pick the BIR band and compute adjustment = amount * (bps/10000).
    Returns rows suitable for display.
    """
    rows = []
    for label, total in category_totals.items():
        bps = find_bir_for_amount(total, bir_tiers)
        if bps is None:
            adj = None
            rate = None
        else:
            rate = bps / 10000.0
            adj = total * rate
        rows.append({
            "category": label,
            "total_amount": total,
            "bps": bps,
            "rate": rate,
            "adjustment": adj
        })
    return rows

def fmt_currency(v: Optional[float]) -> str:
    if v is None:
        return "—"
    return "${:,.2f}".format(v)

def fmt_bps(v: Optional[float]) -> str:
    if v is None:
        return "—"
    return f"{v:.1f} bps"

# -------------------------
# Flask App
# -------------------------
app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev-secret-key")

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        calculation_type = request.form.get("calculation_type", "cv_tier")
        excel = request.files.get("excel_file")
        contract = request.files.get("contract_file")

        if not excel or not contract:
            flash("Please upload both the Excel calculation file and the contract Word file.", "error")
            return redirect(url_for("index"))

        if not is_excel(excel.filename):
            flash("The first file must be an Excel file (.xlsx, .xls, .xlsm, .xlsb).", "error")
            return redirect(url_for("index"))
        if not is_word(contract.filename):
            flash("The second file must be a Word file (.docx).", "error")
            return redirect(url_for("index"))

        upload_dir = os.path.join("uploads")
        os.makedirs(upload_dir, exist_ok=True)
        excel_path = os.path.join(upload_dir, secure_filename(excel.filename))
        contract_path = os.path.join(upload_dir, secure_filename(contract.filename))
        excel.save(excel_path)
        contract.save(contract_path)

        try:
            if calculation_type == "cv_tier":
                return handle_cv_tier_calculation(excel_path, contract_path, request.form)
            else:  # nacv_based
                return handle_nacv_calculation(excel_path, contract_path, request.form)
        except Exception as e:
            flash(f"Error: {e}", "error")
            return redirect(url_for("index"))

    return render_template("index.html")

def handle_cv_tier_calculation(excel_path: str, contract_path: str, form_data):
    """Handle CV Tier-based calculations"""
    # 1) Read Excel and compute totals
    df = read_excel_flex(excel_path)
    category_totals = compute_cv_tier_totals(df)

    if not category_totals:
        flash("No 'CV Tier' or 'Lo-ROC' rows were detected in the Excel file. Please check the format.", "warning")

    # 2) Extract BIR table from Word
    bir_tiers = extract_bir_table_from_contract(contract_path)
    if not bir_tiers:
        flash("No BIR table found in the contract. Ensure the contract contains a NACV/CC-NACV vs BIR bp table.", "warning")

    # 3) Extract CHD Benchmark from Contract
    chd_benchmark = extract_chd_benchmark_from_contract(contract_path)
    chd_input_raw = form_data.get("chd_value")
    chd_input = None
    try:
        chd_input = float(chd_input_raw) if chd_input_raw not in (None, "",) else None
    except:
        flash("Could not parse the CHD input value. Please enter a number like 5.46.", "warning")
        chd_input = None

    chd_diff = (None if (chd_input is None or chd_benchmark is None) else (chd_input - chd_benchmark))

    # 4) Per-category adjustments
    adjustments = calc_per_tier_adjustments(category_totals, bir_tiers)
    total_gross_incentive = sum((r["adjustment"] or 0.0) for r in adjustments)

    # Prepare display versions
    bir_view = []
    for t in bir_tiers:
        lo = f"{t['lower_m']:.2f}M"
        hi = "∞" if t["upper_m"] is None else f"{t['upper_m']:.2f}M"
        bir_view.append({"range": f"{lo} – {hi}", "bps": f"{t['bps']:.1f}"})

    # Build tables for template
    chd_benchmark_display = f"{chd_benchmark:.2f}" if chd_benchmark is not None else "—"
    chd_input_display = f"{chd_input:.2f}" if chd_input is not None else "—"
    chd_diff_display = (f"{chd_diff:.2f}" if chd_diff is not None else "—")
    totals_view = [{"category": k, "total": fmt_currency(v)} for k, v in category_totals.items()]
    adj_view = [{
        "category": r["category"],
        "total": fmt_currency(r["total_amount"]),
        "bps": fmt_bps(r["bps"]),
        "rate": f"{(r['rate']*100):.4f}%" if r["rate"] is not None else "—",
        "adjustment": fmt_currency(r["adjustment"]),
    } for r in adjustments]

    return render_template(
        "result_cv_tier.html",
        chd_benchmark_display=chd_benchmark_display,
        chd_input_display=chd_input_display,
        chd_diff_display=chd_diff_display,
        totals_view=totals_view,
        bir_view=bir_view,
        adj_view=adj_view,
        total_gross=fmt_currency(total_gross_incentive)
    )

def handle_nacv_calculation(excel_path: str, contract_path: str, form_data):
    """Handle NACV-based calculations"""
    # Set currency formatting to GBP for NACV calculations
    format_currency._use_gbp = True
    
    # Calculate monthly totals
    totals = calculate_monthly_totals(excel_path)
    
    # Extract tiers and CHD rules from Word contract
    tiers = extract_tiers_from_docx(contract_path)
    chd_rules = extract_chd_rules_from_docx(contract_path)
    
    # Calculate gross incentive
    gross, breakdown = compute_gross_incentive(totals['nacv_total'], tiers)
    
    # Apply CHD adjustment if CHD performance is provided
    chd_adjustment = 0.0
    chd_breakdown = {"applied": False}
    net_incentive = gross
    
    chd_performance = form_data.get("chd_value", "").strip()
    if chd_performance:
        try:
            chd_days = float(chd_performance)
            chd_adjustment, chd_breakdown = compute_chd_adjustment(gross, chd_days, chd_rules)
            
            # Apply adjustment based on type (deduct or add)
            if chd_breakdown.get("adjustment_type") == "add":
                net_incentive = gross + chd_adjustment
            else:  # Default to deduct
                net_incentive = gross - chd_adjustment
                
        except ValueError:
            flash("Invalid CHD performance value. Please enter a number.", "error")
    
    # Build tiers view for display
    tiers_view = []
    for t in tiers:
        lo = t["lower_m"]
        hi = t["upper_m"]
        bps = t["bps"]
        if hi is None:
            tiers_view.append(f"{lo:.2f}M+ : {bps:.1f} bps")
        else:
            tiers_view.append(f"{lo:.2f}M – {hi:.2f}M : {bps:.1f} bps")

    return render_template(
        "result_nacv.html",
        tiers=tiers_view,
        nacv_value=totals['nacv_total'],
        gross_value=gross,
        breakdown=breakdown,
        chd_rules=chd_rules,
        chd_adjustment=chd_adjustment,
        chd_breakdown=chd_breakdown,
        net_incentive=net_incentive,
        total_bus=totals['bus_total'],
        total_losses=totals['total_losses'],
        bus_col="Bus Unadjusted wo cash - Client",
        ar_col="AR 180 Days Loss", 
        wr_col="Writeoffs - New Reserve Debits",
        format_currency=format_currency_gbp
    )


@app.route("/analyze-files", methods=["POST"])
def analyze_files():
    """AJAX endpoint to analyze files and return auto-detection results"""
    try:
        excel = request.files.get("excel_file")
        contract = request.files.get("contract_file")
        
        if not excel or not contract:
            return {"error": "Both files are required for analysis"}, 400
            
        if not is_excel(excel.filename) or not is_word(contract.filename):
            return {"error": "Invalid file types"}, 400
            
        # Save files temporarily for analysis
        upload_dir = os.path.join("uploads", "temp")
        os.makedirs(upload_dir, exist_ok=True)
        excel_path = os.path.join(upload_dir, "temp_" + secure_filename(excel.filename))
        contract_path = os.path.join(upload_dir, "temp_" + secure_filename(contract.filename))
        excel.save(excel_path)
        contract.save(contract_path)
        
        # Perform auto-detection
        detection_result = detect_calculation_type(excel_path)
        contract_validation = validate_detection_with_contract(contract_path, detection_result['detected_type'])
        
        # Adjust confidence based on contract validation
        final_confidence = detection_result['confidence']
        if contract_validation['confidence_boost'] == 1:
            if final_confidence == 'medium':
                final_confidence = 'high'
            elif final_confidence == 'low':
                final_confidence = 'medium'
        elif contract_validation['confidence_boost'] == -1:
            if final_confidence == 'high':
                final_confidence = 'medium'
            elif final_confidence == 'medium':
                final_confidence = 'low'
        
        # Clean up temp files
        try:
            os.remove(excel_path)
            os.remove(contract_path)
        except:
            pass
            
        return {
            "detection": {
                **detection_result,
                "confidence": final_confidence
            },
            "contract_validation": contract_validation,
            "success": True
        }
        
    except Exception as e:
        return {"error": str(e)}, 500

@app.route("/healthz")
def healthz():
    return {"ok": True}


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
